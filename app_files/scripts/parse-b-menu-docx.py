#!/usr/bin/env python3
"""Extract Camp Walden B Menu offerings from a DOCX table.

Usage:
  python scripts/parse-b-menu-docx.py "/path/to/B MENU 2023 S2.docx"

The parser preserves the camp-specific rules:
- "/18 Water-skiing" becomes rosterLimit 18.
- "(1,2)" becomes eligibleUnits ["UNIT1", "UNIT2"].
- entries without numeric limits become SPECIAL_APPROVAL.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


PERIODS = ["P1B", "P2B", "P3B", "P4B"]
UNIT_MAP = {"1": "UNIT1", "2": "UNIT2", "3": "UNIT3", "4": "UNIT4"}
ALL_UNITS = ["UNIT1", "UNIT2", "UNIT3", "UNIT4"]
ALL_SWIM = ["BLUEGILL", "WALLEYE", "MUSKIE"]


def extract_cell_lines(cell_text: str) -> list[str]:
    return [line.strip() for line in cell_text.splitlines() if line.strip()]


def parse_entry(period: str, area: str, entry: str) -> dict | None:
    if not entry or entry == "/":
        return None

    match = re.match(r"^/\s*(?:(\d+)\s*)?(.*)$", entry.strip())
    if not match:
        return None

    limit_raw, name_raw = match.groups()
    name = re.sub(r"\s+", " ", name_raw).strip()
    if not name:
        return None

    units = ALL_UNITS
    unit_match = re.search(r"\(([\d,\s]+)\)", name)
    if unit_match:
        units = [UNIT_MAP[number.strip()] for number in unit_match.group(1).split(",") if number.strip() in UNIT_MAP]
        name = re.sub(r"\s*\([\d,\s]+\)", "", name).strip()

    name = name.replace("Ski", "Water-skiing") if name == "Ski" else name
    name = name.replace("Water-skiing \u2013 All levels", "Water-skiing")

    return {
        "period": period,
        "area": area,
        "activity": name,
        "rosterLimit": int(limit_raw) if limit_raw else None,
        "limitType": "FIXED" if limit_raw else "SPECIAL_APPROVAL",
        "eligibleUnits": units,
        "eligibleSwimLevels": ["BLUEGILL"] if name == "Blue Gill Swim" else ALL_SWIM,
        "preAssigned": area.upper() == "RIDING",
        "notes": "Imported from B menu"
    }


def parse_docx(path: Path) -> list[dict]:
    doc = Document(path)
    if len(doc.tables) < 2:
        return []

    table = doc.tables[1]
    offerings: list[dict] = []

    for row in table.rows[1:]:
        for index, cell in enumerate(row.cells):
            period = PERIODS[index]
            area = ""
            for line in extract_cell_lines(cell.text):
                if not line.startswith("/"):
                    area = line.strip()
                    continue
                parsed = parse_entry(period, area, line)
                if parsed:
                    offerings.append(parsed)

    return offerings


def main() -> int:
    if len(sys.argv) != 2:
        print(__doc__)
        return 2

    offerings = parse_docx(Path(sys.argv[1]))
    print(json.dumps(offerings, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
